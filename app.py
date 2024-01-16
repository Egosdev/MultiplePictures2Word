import os
from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

directory = "D:\5"

def get_jpg_files(directory):
    jpg_files = [file for file in os.listdir(directory) if file.lower().endswith('.jpg')]
    return jpg_files

def set_cell_margins(cell, margin_value):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run_element = run._element
            cell_element = run_element.getparent()
            cell_element_margins = cell_element.find('.//w:tblCellMar', namespaces=cell_element.nsmap)

            if cell_element_margins is None:
                cell_element_margins = OxmlElement('w:tblCellMar')
                cell_element.append(cell_element_margins)

            cell_element_margins.clear_content()

            for side in ('top', 'start', 'bottom', 'end'):
                margin = OxmlElement(f'w:{side}')
                margin.set(qn('w:w'), f'{margin_value}cm')
                margin.set(qn('w:type'), 'dxa')
                cell_element_margins.append(margin)

def create_word_table(doc, jpg_files):
    table_rows = len(jpg_files) // 2 + 1
    table_cols = 2
    
    table = doc.add_table(rows=table_rows, cols=table_cols)
    table.autofit = False

    # Set cell margins to reduce the gap
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, 0)  # Set the margin value as needed

    for i in range(table_rows):
        for j in range(table_cols):
            index = i * table_cols + j
            if index < len(jpg_files):
                cell = table.cell(i, j)
                img_path = os.path.join(directory, jpg_files[index])
                cell.add_paragraph().add_run().add_picture(img_path, height=Cm(6.17), width=Cm(8.23))


def add_paragraph_with_style(doc, text, font_name, font_size, alignment):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)

    font = run.font
    font.name = font_name
    font.size = Pt(font_size)

    paragraph.alignment = alignment

def main():
    doc = Document()
    for i in range(1, 58):
        global directory
        directory = f"D:\{i}"
        try: 

            jpg_files = get_jpg_files(directory)

            if jpg_files:
                add_paragraph_with_style(doc, f"{i} NOLU BLOK", "Times New Roman", 12, WD_PARAGRAPH_ALIGNMENT.LEFT)
                create_word_table(doc, jpg_files)

                # Add a blank page between directories
                doc.add_page_break()
                print(i, "done")
        except FileNotFoundError:
            print("not found")
            continue
    doc.save(f"output.docx")


if __name__ == "__main__":
    main()
